import pandas as pd
import seaborn as sns
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

def create_report(sales: pd.DataFrame, output_path: str = "report.docx"):
    sns.set(style="whitegrid", palette="pastel")

    report = Document()
    report.add_heading("ClickHouse Sales Report", level=1)
    report.add_paragraph("The following report presents basic analysis of diamond sales based on data from ClickHouse.")
    
    plt.figure(figsize=(8, 5))
    sns.histplot(sales, x="total_value", bins=30, kde=True)
    plt.title("Distribution of Total Sales Value")
    plt.tight_layout()
    plt.savefig("plot1_total_value.png")
    plt.close()

    report.add_heading("Distribution of Total Sales Value", level=2)
    report.add_picture("plot1_total_value.png", width=Inches(6))
    report.add_paragraph("The histogram shows the distribution of sales value for all transactions.")

    categories = ["cut", "color", "clarity"]
    
    for cat in categories:
        report.add_heading(f"Analysis by {cat.capitalize()}", level=2)
        grouped = sales.groupby(cat)["total_value"].agg(["count", "mean", "sum"]).reset_index()

        for metric in ["count", "mean", "sum"]:
            plt.figure(figsize=(8, 5))
            sns.barplot(data=grouped, x=cat, y=metric)
            plt.title(f"{metric.capitalize()} by {cat.capitalize()}")
            plt.tight_layout()
            filename = f"plot_{cat}_{metric}.png"
            plt.savefig(filename)
            plt.close()

            report.add_heading(f"{metric.capitalize()} by {cat.capitalize()}", level=3)
            report.add_picture(filename, width=Inches(6))
            report.add_paragraph(f"The chart displays the {metric} of sales value broken down by {cat}.")

    report.save(output_path)

    for file in os.listdir():
        if file.endswith(".png") and file.startswith(("plot1", "plot_")):
            os.remove(file)

    print(f"Report successfully saved as: {output_path}")